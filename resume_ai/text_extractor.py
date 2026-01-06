# import pdfplumber
# import docx
# import os

# def extract_text(file_path):
#     if file_path.endswith(".pdf"):
#         with pdfplumber.open(file_path) as pdf:
#             return "\n".join(
#                 page.extract_text() or "" for page in pdf.pages
#             )

#     if file_path.endswith(".docx"):
#         doc = docx.Document(file_path)
#         return "\n".join(p.text for p in doc.paragraphs)

#     return ""


import pdfplumber
import docx


def extract_text(file_path):
    """
    Extract text from PDF or DOCX resume
    """
    text = ""

    if file_path.lower().endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"

    elif file_path.lower().endswith(".docx"):
        doc = docx.Document(file_path)
        for p in doc.paragraphs:
            text += p.text + "\n"

    return text
